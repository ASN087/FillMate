from django.shortcuts import render, get_object_or_404
from rest_framework import status
from rest_framework.generics import ListCreateAPIView, RetrieveAPIView, ListAPIView
from rest_framework.permissions import IsAdminUser, IsAuthenticated
from .permissions import IsHODUser
from rest_framework.response import Response
from rest_framework.views import APIView
from .models import DocumentTemplate, Placeholder, SubmittedDocument, GeneratedDocument
from .serializers import DocumentTemplateSerializer, PlaceholderSerializer, SubmittedDocumentSerializer, DocumentReviewSerializer
from .utils import extract_placeholders_from_docx, convert_docx_to_pdf, generate_signed_pdf
from django.http import HttpResponse, JsonResponse, FileResponse
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from io import BytesIO
import re
import tempfile
import json
import os
import logging
from django.core.files.base import ContentFile
from documents.models import GeneratedDocument, ApprovedDocument
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import Group
from django.urls import reverse
from users.models import UserProfile  # Replace 'your_app' with the actual app name where UserProfile is defined
from notifications.models import Notification
from django.db.models.signals import post_save
from django.dispatch import receiver
from rest_framework.exceptions import ValidationError
from django.contrib.contenttypes.models import ContentType
from django.views.decorators.clickjacking import xframe_options_exempt
from notifications.utils import notify_document_submission




logger = logging.getLogger(__name__)

# ✅ API to List All Document Templates (For All Users) & Upload Templates (Only Admins)
class DocumentTemplateListCreateView(ListCreateAPIView):
    """API to list all document templates (for authenticated users) & upload templates (only for admins)"""
    
    queryset = DocumentTemplate.objects.all()
    serializer_class = DocumentTemplateSerializer

    def get_permissions(self):
        """Allow normal users to list templates, but restrict uploads to admins."""
        if self.request.method == 'POST':
            return [IsAdminUser()]  # Only admins can upload templates
        return [IsAuthenticated()]  # Normal users can list templates

    def perform_create(self, serializer):
        """Extract placeholders after saving the template (Admin Only)"""
        template = serializer.save()  # Save uploaded template to DB
        extract_placeholders_from_docx(template)  # Extract placeholders

# ✅ API to Fetch a Single Template & Its Placeholders
class DocumentTemplateDetailView(RetrieveAPIView):
    """API to fetch a document template and its placeholders"""
    
    queryset = DocumentTemplate.objects.all()
    serializer_class = DocumentTemplateSerializer
    permission_classes = [IsAuthenticated]  # Only logged-in users can access

# ✅ API to List Placeholders for a Selected Template
class PlaceholderListView(ListAPIView):
    """API to return placeholders based on the selected template ID"""
    
    serializer_class = PlaceholderSerializer
    permission_classes = [IsAuthenticated]  # Users must be logged in

    def get_queryset(self):
        """Return placeholders for a specific template"""
        template_id = self.kwargs['template_id']
        return Placeholder.objects.filter(template_id=template_id)

def preview_template(request, template_id):
    template = get_object_or_404(DocumentTemplate, pk=template_id)
    
    try:
        # Debugging - log that we're starting preview generation
        logger.info(f"Starting preview generation for template: {template.name}")
        
        # Create PDF buffer
        buffer = BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        story.append(Paragraph(f"Template Preview: {template.name}", styles['Title']))
        story.append(Spacer(1, 12))
        
        try:
            # Process DOCX content
            word_doc = Document(BytesIO(template.file.read()))
            
            # Process paragraphs
            for para in word_doc.paragraphs:
                if para.text.strip():
                    # Highlight placeholders
                    text = re.sub(r'<([^>]+)>', r'<font color="orange">&lt;\1&gt;</font>', para.text)
                    p = Paragraph(text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 8))
            
            # Process tables
            for table in word_doc.tables:
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = re.sub(r'<([^>]+)>', r'<font color="orange">&lt;\1&gt;</font>', cell.text)
                        row_data.append(Paragraph(cell_text, styles['Normal']))
                    data.append(row_data)
                
                tbl = Table(data)
                tbl.setStyle(TableStyle([
                    ('GRID', (0,0), (-1,-1), 1, colors.grey),
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            # Return PDF response
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'inline; filename="{template.name}_preview.pdf"'
            logger.info("Successfully generated PDF preview")
            return response
            
        except Exception as e:
            logger.error(f"Error processing DOCX content: {str(e)}", exc_info=True)
            return JsonResponse({
                'error': 'Failed to process template content',
                'details': str(e)
            }, status=500)
            
    except Exception as e:
        logger.error(f"Preview generation failed: {str(e)}", exc_info=True)
        return JsonResponse({
            'error': 'Failed to generate preview',
            'details': str(e)
        }, status=500)
    
def clean_placeholder(text):
    """
    Cleans placeholders by removing example text inside parentheses.
    Example: 
    "<OFFICER_DESIGNATION (e.g., SHO)>" → "<OFFICER_DESIGNATION>"
    "<ISSUING_AUTHORITY (e.g., Sub Divisional Magistrate)>" → "<ISSUING_AUTHORITY>"
    "<OFFENCE_SECTION (e.g., U/s 126/129 BNSS)>" → "<OFFENCE_SECTION>"
    """
    return re.sub(r'(<[A-Z_]+)\s*\(.*?\)>', r'\1>', text)

def extract_placeholders(doc):
    """
    Extracts all placeholders from the document text, including placeholders with apostrophes.
    """
    placeholders_found = set()
    
    # Updated regex to include apostrophes inside placeholders
    placeholder_pattern = r'<[A-Z_\'-]+(?:\s*\(.*?\))?>'

    for para in doc.paragraphs:
        matches = re.findall(placeholder_pattern, para.text)
        placeholders_found.update(matches)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(placeholder_pattern, cell.text)
                placeholders_found.update(matches)

    return placeholders_found

def generate_document(request, template_id):
    """Generate a document by replacing placeholders with user-provided values."""
    template = get_object_or_404(DocumentTemplate, pk=template_id)

    try:
        # 1. Load template
        doc = Document(BytesIO(template.file.read()))

        # 2. Get all placeholders from database (Map to user inputs)
        db_placeholders = {p.placeholder_text: p.name for p in template.placeholders.all()}

        # 3. Extract placeholders from document and clean them
        doc_placeholders = {ph: clean_placeholder(ph) for ph in extract_placeholders(doc)}

        # 🔍 DEBUG: Show extracted & cleaned placeholders
        print("📌 Placeholders in Document (Raw):", list(doc_placeholders.keys()))
        print("📌 Cleaned Placeholders for Matching:", list(doc_placeholders.values()))
        print("📌 Database Placeholders:", list(db_placeholders.keys()))

        # 4. Replace placeholders
        replacements_made = 0

        def replace_placeholders_in_text(text):
            modified_text = text
            for raw_placeholder, cleaned_placeholder in doc_placeholders.items():
                if cleaned_placeholder in db_placeholders:
                    field_name = db_placeholders[cleaned_placeholder]
                    user_value = request.POST.get(field_name, "").strip()
                    if user_value:
                        modified_text = modified_text.replace(raw_placeholder, user_value)
            return modified_text

        # Replace in paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            para.text = replace_placeholders_in_text(para.text)
            if para.text != original_text:
                replacements_made += 1

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    cell.text = replace_placeholders_in_text(cell.text)
                    if cell.text != original_text:
                        replacements_made += 1

        # 🔍 DEBUG: Check if replacements happened
        print(f"✅ Total Replacements Made: {replacements_made}")
        if replacements_made == 0:
            raise ValueError("No placeholders were replaced! Check if document placeholders match database.")

        # 5. Save and return the document
        output_format = request.POST.get('format', 'docx')
        buffer = BytesIO()

        if output_format == 'pdf':
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            pdf_buffer = convert_docx_to_pdf(docx_buffer)
            buffer = pdf_buffer
            content_type = 'application/pdf'
            file_extension = 'pdf'
        else:
            doc.save(buffer)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_extension = 'docx'

        # Ensure buffer is at the beginning before saving
        buffer.seek(0)

        # Save generated document in database
        generated_doc = GeneratedDocument(user=request.user)
        file_name = f"{template.name}_{request.user.first_name}.{file_extension}"
        generated_doc.file.save(file_name, ContentFile(buffer.read()))
        generated_doc.save()  # Save the record in the database

        # Return response
        return HttpResponse(
            buffer.getvalue(),
            content_type=content_type,
            headers={'Content-Disposition': f'attachment; filename="{template.name}.{file_extension}"'}
        )

    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


@login_required # Ensures only logged-in users can access
def my_documents(request):
    """
    Displays documents submitted by the currently logged-in user.
    Includes status and download links for approved documents.
    """
    # Filter SubmittedDocument for the current user
    # Order by most recent submission first
    # Use select_related for template (one-to-one/foreignkey)
    # Use prefetch_related for approved_version (one-to-one reverse)
    user_submissions = SubmittedDocument.objects.filter(
        user=request.user
    ).select_related(
        'template' # To get template.name efficiently
    ).prefetch_related(
        'approved_version' # To get the linked ApprovedDocument efficiently
    ).order_by('-submitted_at') # Show newest first

    context = {
        'submissions': user_submissions,
        'page_title': 'My Submitted Documents' # Optional: for template title
    }
    # Ensure you have a template named 'my_documents.html'
    return render(request, 'my_documents.html', context)



class SubmitDocumentView(APIView):

    
    permission_classes = [IsAuthenticated]

    def post(self, request, template_id):
        if not DocumentTemplate.objects.filter(id=template_id).exists():
            return Response({"error": "Invalid template ID"}, status=status.HTTP_400_BAD_REQUEST)
        try:
            template = get_object_or_404(DocumentTemplate, id=template_id)

            # Verify CSRF token
            if not request.META.get('HTTP_X_CSRFTOKEN') == request.COOKIES.get('csrftoken'):
                return Response({"error": "CSRF verification failed"}, status=status.HTTP_403_FORBIDDEN)
            
            # Generate document
            doc = Document(BytesIO(template.file.read()))
            
            # Get all placeholders from database
            db_placeholders = {p.placeholder_text: p.name for p in template.placeholders.all()}
            doc_placeholders = {ph: clean_placeholder(ph) for ph in extract_placeholders(doc)}

            # Handle both form-data and JSON input
            post_data = request.POST if request.POST else json.loads(request.body)
            
            # Replace placeholders with form data
            for para in doc.paragraphs:
                para.text = replace_placeholders_in_text(para.text, doc_placeholders, db_placeholders, request.POST)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = replace_placeholders_in_text(cell.text, doc_placeholders, db_placeholders, request.POST)
            
            # Save to buffer based on format
            output_format = request.POST.get('format', 'docx')
            buffer = BytesIO()
            
            if output_format == 'pdf':
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                pdf_buffer = convert_docx_to_pdf(docx_buffer)
                buffer = pdf_buffer
                file_extension = 'pdf'
                content_type = 'application/pdf'
            else:
                doc.save(buffer)
                file_extension = 'docx'
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            buffer.seek(0)
            
            # Create SubmittedDocument record
            submitted_doc = SubmittedDocument.objects.create(
                user=request.user,
                template=template,
                status='Pending'
            )
            
            # Save the generated file
            file_name = f"submitted_{template.name}_{request.user.username}.{file_extension}"
            submitted_doc.document.save(file_name, ContentFile(buffer.read()))

            # ✅ REPLACE WITH THIS CALL to the utility function:
            try:
                notify_document_submission(submitted_doc, request.user)
                logger.info(f"HOD notification process initiated for submission {submitted_doc.id}")
            except Exception as notify_error:
                logger.error(f"Failed to initiate HOD notifications for submission {submitted_doc.id}: {notify_error}", exc_info=True)
            # ✅ --- END HOD NOTIFICATION --- ✅
            
            return Response({
                "status": "success",
                "submission_id": submitted_doc.id,
                "redirect_url": reverse('documents:my_documents') # Redirect to the actual page URL
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            # Log error details
            logger.error(f"Error during document submission by {request.user.username} for template {template_id}: {e}", exc_info=True)
            # Provide a generic error message
            return Response({
                "status": "error",
                "error": f"An error occurred during submission: {e}"
            }, status=status.HTTP_400_BAD_REQUEST)

# Helper function
def replace_placeholders_in_text(text, doc_placeholders, db_placeholders, post_data):
    for raw_placeholder, cleaned_placeholder in doc_placeholders.items():
        if cleaned_placeholder in db_placeholders:
            field_name = db_placeholders[cleaned_placeholder]
            user_value = post_data.get(field_name, "").strip()
            if user_value:
                text = text.replace(raw_placeholder, user_value)
    return text

class SubmissionDetailView(RetrieveAPIView):
    queryset = SubmittedDocument.objects.all()
    serializer_class = SubmittedDocumentSerializer
    permission_classes = [IsAuthenticated]
    lookup_field = 'submission_id'  # Use 'submission_id' to match the URL pattern

    def get_queryset(self):
        # Only allow HODs or the original submitter to view
        if self.request.user.groups.filter(name='HOD').exists():
            return super().get_queryset()
        return super().get_queryset().filter(user=self.request.user)
    

# documents/views.py
class DocumentReviewView(APIView):
    permission_classes = [IsAuthenticated, IsHODUser]  # Create IsHODUser permission
    @xframe_options_exempt
    def get(self, request, submission_id):
        submission = get_object_or_404(SubmittedDocument, id=submission_id)
        # --- START: Suggestion for GET logic ---
        # Ensure submitted doc exists
        if not submission.document:
             return Response({'error': 'Submitted document file not found.'}, status=404)

        file_path = submission.document.path
        content_type = 'application/octet-stream' # Default
        file_name = submission.document.name

        if file_name.lower().endswith('.pdf'):
             content_type = 'application/pdf'
        elif file_name.lower().endswith('.docx'):
            # Option 1: Convert to PDF on the fly for viewing (can be slow)
            try:
                with open(file_path, 'rb') as docx_file:
                    docx_buffer = BytesIO(docx_file.read())
                pdf_buffer = convert_docx_to_pdf(docx_buffer) # Your existing util function
                response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
                # Optional: Set filename for inline view
                # response['Content-Disposition'] = f'inline; filename="preview_{os.path.basename(file_name)}.pdf"'
                return response
            except Exception as e:
                 logger.error(f"Error converting DOCX to PDF for review: {e}", exc_info=True)
                 return Response({'error': f'Could not convert document to PDF for preview: {e}'}, status=500)
            # Option 2: Tell the frontend it's DOCX (iframe won't work)
            # return Response({'error': 'Preview not available for DOCX. Please download.'}, status=400)
        else:
             # Handle other file types or return error
             return Response({'error': 'Unsupported document type for preview.'}, status=400)

        # If it was already PDF:
        try:
            return FileResponse(open(file_path, 'rb'), content_type=content_type)
        except FileNotFoundError:
            return Response({'error': 'Document file not found on server.'}, status=404)
        # --- END: Suggestion for GET logic ---


    def post(self, request, submission_id):
        submission = get_object_or_404(SubmittedDocument, pk=submission_id)

        if submission.status != 'Pending':
             return Response({'error': f'Submission is already {submission.status}.'}, status=status.HTTP_400_BAD_REQUEST)

        action = request.data.get('action')
        if not action:
            return Response({'error': 'Missing action (approve/reject)'}, status=status.HTTP_400_BAD_REQUEST)

        if action == 'approve':
            return self._handle_approval(request, submission)
        elif action == 'reject':
            return self._handle_rejection(request, submission)
        else:
            return Response({'error': 'Invalid action'}, status=status.HTTP_400_BAD_REQUEST)

    def _handle_approval(self, request, submission):
        # ✅ Get or create the profile for the HOD user
        profile, created = UserProfile.objects.get_or_create(user=request.user)

        # Check if the signature exists *on the retrieved profile*
        if not profile.digital_signature:
            # If the profile was just created OR if it exists but has no signature
            error_msg = "Approval failed: Please upload your digital signature first via the 'Upload Signature' link."
            # Use DRF's validation error for consistency in API responses
            # return Response({'error': error_msg}, status=status.HTTP_400_BAD_REQUEST)
            raise ValidationError(error_msg) # This is often better practice in DRF views

        # Ensure the submitted document file actually exists on disk
        if not submission.document or not os.path.exists(submission.document.path):
             logger.error(f"Approval failed: Submitted document file missing for submission {submission.id} at path {submission.document.path}")
             raise ValidationError("Approval failed: The original submitted document file cannot be found.")

        # Ensure the signature file actually exists on disk
        if not os.path.exists(profile.digital_signature.path):
            logger.error(f"Approval failed: Signature file missing for user {request.user.username} at path {profile.digital_signature.path}")
            raise ValidationError("Approval failed: Your signature file cannot be found. Please re-upload it.")


        try:
            # --- Determine input path and convert if necessary ---
            input_path_for_signing = submission.document.path
            temp_pdf_buffer = None # To hold converted PDF if needed

            if not submission.document.name.lower().endswith('.pdf'):
                 # Convert DOCX to PDF first
                 logger.info(f"Converting DOCX to PDF for signing: {submission.document.path}")
                 with open(submission.document.path, 'rb') as docx_file:
                     docx_buffer = BytesIO(docx_file.read())
                 temp_pdf_buffer = convert_docx_to_pdf(docx_buffer) # Your util function
                 # We need to pass the *content* of the PDF buffer to generate_signed_pdf,
                 # or save it temporarily and pass the path. Passing buffer is better.
                 # Let's modify generate_signed_pdf slightly if needed, or use the buffer directly.
                 # Assuming generate_signed_pdf can take a file-like object (BytesIO) for the original PDF:
                 input_path_for_signing = temp_pdf_buffer # Use the buffer directly
                 input_path_for_signing.seek(0) # Reset buffer position
                 logger.info("Conversion successful.")
            # --- End Conversion ---


            signed_pdf_buffer = generate_signed_pdf(
                input_path_for_signing, # This is now either the original PDF path or the BytesIO buffer of the converted PDF
                profile.digital_signature.path
            )

            # Create the ApprovedDocument record
            approved_doc = ApprovedDocument.objects.create(
                original_submission=submission,
                approved_by=request.user
                # signed_file will be saved next
            )

            # Define the name for the signed file (ensure it's .pdf)
            base_name = os.path.splitext(os.path.basename(submission.document.name))[0]
            signed_file_name = f'signed_{base_name}.pdf' # Always PDF now

            # Save the signed PDF buffer to the ApprovedDocument
            approved_doc.signed_file.save(
                 signed_file_name,
                 ContentFile(signed_pdf_buffer.read()),
                 save=True # Save the model instance after file save
            )

            # Update the original submission status
            submission.status = 'Approved'
            submission.rejection_reason = None # Clear any previous rejection reason
            submission.save()

            # ✅ --- CREATE NOTIFICATION FOR SUBMITTING USER --- ✅
            try:
                content_type = ContentType.objects.get_for_model(SubmittedDocument)
                submitting_user = submission.user # Get the user who submitted
                hod_user = request.user # The HOD performing the action

                Notification.objects.create(
                    recipient=submitting_user, # <<< Notify the original user
                    sender=hod_user,
                    message=f"Your submission '{submission.template.name}' (ID: {submission.id}) has been approved.",
                    content_type=content_type,
                    object_id=submission.id,
                    is_read=False # Ensure it starts as unread
                )
                logger.info(f"Approval notification created for user {submitting_user.username} for submission {submission.id}")
            except Exception as notify_error:
                 logger.error(f"Failed to create approval notification for submission {submission.id}: {notify_error}", exc_info=True)
            # ✅ --- END NOTIFICATION --- ✅

            logger.info(f"Document {submission.id} approved successfully by {request.user.username}")
            return Response({'status': 'approved', 'message': 'Document approved and signed.'}, status=status.HTTP_200_OK)

        except ValidationError as ve: # Catch validation errors specifically
             logger.warning(f"Approval validation failed for submission {submission.id}: {ve.detail}")
             return Response({'error': ve.detail}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            # Log the full error for debugging
            logger.error(f"Unexpected error during approval of submission {submission.id} by {request.user.username}: {e}", exc_info=True)
            # Provide a generic error to the user
            return Response({'error': f'An unexpected error occurred during approval: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
             # Clean up temporary buffer if created
             if temp_pdf_buffer:
                  temp_pdf_buffer.close()
        
        

    def _handle_rejection(self, request, submission):
        reason = request.data.get('reason', '').strip()
        if not reason:
            # return Response({'error': 'Rejection reason is required'}, status=status.HTTP_400_BAD_REQUEST)
             raise ValidationError("Rejection reason is required.") # Use ValidationError

        try:
            submission.status = 'Rejected'
            submission.rejection_reason = reason # Save the reason
            submission.save()

            # ✅ --- CREATE NOTIFICATION FOR SUBMITTING USER --- ✅
            try:
                content_type = ContentType.objects.get_for_model(SubmittedDocument)
                submitting_user = submission.user # Get the user who submitted
                hod_user = request.user # The HOD performing the action

                # Truncate reason for notification message if too long
                truncated_reason = (reason[:75] + '...') if len(reason) > 75 else reason

                Notification.objects.create(
                    recipient=submitting_user, # <<< Notify the original user
                    sender=hod_user,
                    message=f"Your submission '{submission.template.name}' (ID: {submission.id}) was rejected. Reason: {truncated_reason}",
                    content_type=content_type,
                    object_id=submission.id,
                    is_read=False # Ensure it starts as unread
                )
                logger.info(f"Rejection notification created for user {submitting_user.username} for submission {submission.id}")
            except Exception as notify_error:
                 logger.error(f"Failed to create rejection notification for submission {submission.id}: {notify_error}", exc_info=True)
            # ✅ --- END NOTIFICATION --- ✅

            logger.info(f"Document {submission.id} rejected by {request.user.username}. Reason: {reason}")

            # TODO: Optionally send a notification back to the submitting user about the rejection

            return Response({'status': 'rejected', 'message': 'Document rejected.'}, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"Error during rejection of submission {submission.id} by {request.user.username}: {e}", exc_info=True)
            return Response({'error': f'An unexpected error occurred during rejection: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
